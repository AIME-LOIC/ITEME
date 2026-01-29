# ...existing code...
import os
import csv
import io
import base64
import json
from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from config import supabase  # ✅ Import Supabase client

app = FastAPI()

# Admin password - set via environment variable or use default
ADMIN_PASSWORD = os.getenv('ADMIN_PASSWORD', 'iteme2026')

# --- API ROUTES ---

@app.get("/")
def home(request: Request):
    """Serves the public submission form."""
    return HTMLResponse(PUBLIC_HTML_TEMPLATE)

@app.post('/api/submit')
async def submit_data(
    name: str = Form(...),
    date: str = Form(...),
    amount: str = Form(...),
    receiver: str = Form(...)
):
    """Endpoint to save form data to Supabase."""
    try:
        data = {
            "name": name,
            "date": date,
            "amount": float(amount),
            "receiver": receiver
        }
        
        result = supabase.table("payments").insert(data).execute()
        
        return JSONResponse(
            {"status": "success", "message": "Data saved successfully!"},
            status_code=201
        )
    except Exception as e:
        return JSONResponse(
            {"status": "error", "message": str(e)},
            status_code=500
        )

@app.get('/api/admin/data')
async def get_data(pwd: str = None):
    """Endpoint for the admin panel to fetch all records."""
    if pwd:
        try:
            pwd = base64.b64decode(pwd).decode('utf-8')
        except:
            pass
    if not pwd or pwd != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    try:
        result = supabase.table("payments").select("*").order("created_at", desc=True).execute()
        return result.data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Single delete endpoint (admin only).
@app.delete('/api/admin/delete')
async def delete_payment(id: int = None, pwd: str = None):
    """Delete a single payment record by id (admin only)."""
    if pwd:
        try:
            pwd = base64.b64decode(pwd).decode('utf-8')
        except:
            pass
    if not pwd or pwd != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Unauthorized")
    if id is None:
        raise HTTPException(status_code=400, detail="Missing id parameter")
    try:
        result = supabase.table("payments").delete().eq("id", id).execute()
        # Supabase may return deleted rows in result.data; if none, treat as not found
        if hasattr(result, "data") and (result.data is None or (isinstance(result.data, list) and len(result.data) == 0)):
            raise HTTPException(status_code=404, detail="Payment not found")
        return JSONResponse({"status": "success", "message": "Payment deleted", "deleted": result.data})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# New: bulk delete endpoint (admin only) - accepts JSON body: {"ids": [1,2,3]}
@app.post('/api/admin/delete-multiple')
async def delete_multiple(payload: dict, pwd: str = None):
    """
    Delete multiple payments by IDs.
    Send JSON body: {"ids": [1,2,3]}
    """
    if pwd:
        try:
            pwd = base64.b64decode(pwd).decode('utf-8')
        except:
            pass
    if not pwd or pwd != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Unauthorized")
    ids = payload.get("ids") if isinstance(payload, dict) else None
    if not ids or not isinstance(ids, list):
        raise HTTPException(status_code=400, detail="Missing or invalid 'ids' list in request body")
    try:
        # Supabase client: use .in_('id', ids) if available
        result = supabase.table("payments").delete().in_("id", ids).execute()
        # result.data may contain deleted rows; return them
        return JSONResponse({"status": "success", "deleted": result.data})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get('/api/admin/download')
async def download_csv(pwd: str = None, format: str = 'csv'):
    """Download all transaction data in selected format."""
    if pwd:
        try:
            pwd = base64.b64decode(pwd).decode('utf-8')
        except:
            pass
    if not pwd or pwd != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    try:
        result = supabase.table("payments").select("*").order("created_at", desc=True).execute()
        rows = result.data
        
        if format == 'csv':
            # Create CSV in memory
            output = io.StringIO()
            if rows:
                fieldnames = list(rows[0].keys())
                writer = csv.DictWriter(output, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(rows)
            
            csv_data = output.getvalue()
            output.close()
            
            return FileResponse(
                io.BytesIO(csv_data.encode('utf-8')),
                media_type='text/csv',
                filename='transactions_data.csv'
            )
        
        elif format == 'excel':
            try:
                import openpyxl
                from openpyxl.styles import Font, PatternFill, Alignment
                
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Transactions"
                
                if rows:
                    headers = list(rows[0].keys())
                    for col_idx, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col_idx, value=header)
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill(start_color="007BFF", end_color="007BFF", fill_type="solid")
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                    
                    for row_idx, row_data in enumerate(rows, 2):
                        for col_idx, header in enumerate(headers, 1):
                            ws.cell(row=row_idx, column=col_idx, value=row_data[header])
                
                output = io.BytesIO()
                wb.save(output)
                output.seek(0)
                
                return FileResponse(
                    output,
                    media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    filename='transactions_data.xlsx'
                )
            except ImportError:
                raise HTTPException(status_code=400, detail="Excel support not installed. Please install openpyxl.")
        
        elif format == 'pdf':
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
                from reportlab.lib import colors
                
                output = io.BytesIO()
                doc = SimpleDocTemplate(output, pagesize=letter)
                elements = []
                
                if rows:
                    headers = list(rows[0].keys())
                    data = [headers]
                    for row in rows:
                        data.append([str(row[h]) for h in headers])
                    
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#007BFF')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    elements.append(table)
                
                doc.build(elements)
                output.seek(0)
                
                return FileResponse(
                    output,
                    media_type='application/pdf',
                    filename='transactions_data.pdf'
                )
            except ImportError:
                raise HTTPException(status_code=400, detail="PDF support not installed. Please install reportlab.")
        
        elif format == 'docx':
            try:
                from docx import Document
                
                doc = Document()
                doc.add_heading('Transaction Report', 0)
                
                if rows:
                    headers = list(rows[0].keys())
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    header_row = table.rows[0]
                    for col_idx, header in enumerate(headers):
                        cell = header_row.cells[col_idx]
                        cell.text = header
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for col_idx, header in enumerate(headers):
                            row_cells[col_idx].text = str(row_data[header])
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                return FileResponse(
                    output,
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    filename='transactions_data.docx'
                )
            except ImportError:
                raise HTTPException(status_code=400, detail="DOCX support not installed. Please install python-docx.")
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get('/admin')
async def admin_panel(pwd: str = None):
    """Admin panel - password protected."""
    if pwd:
        try:
            pwd = base64.b64decode(pwd).decode('utf-8')
        except:
            pass
    if not pwd or pwd != ADMIN_PASSWORD:
        return HTMLResponse(ADMIN_LOGIN_TEMPLATE, status_code=401)
    return HTMLResponse(ADMIN_HTML_TEMPLATE)

# --- HTML/CSS/JS ASSETS ---

ADMIN_LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Admin Login</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gradient-to-r from-blue-500 to-blue-700 min-h-screen flex items-center justify-center">
    <div class="bg-white p-8 rounded-lg shadow-2xl w-96">
        <h2 class="text-2xl font-bold mb-6 text-center text-gray-800">Admin Login</h2>
        <form id="loginForm" class="space-y-4">
            <div>
                <label class="block text-sm font-medium text-gray-700 mb-2">Admin Password</label>
                <input type="password" id="pwdInput" required autofocus class="w-full p-2.5 border border-gray-300 rounded-lg focus:ring-2 focus:ring-blue-500 outline-none">
            </div>
            <button type="submit" class="w-full bg-blue-600 text-white font-bold py-2 rounded-lg hover:bg-blue-700 transition"> Login</button>
        </form>
        <p class="text-center text-gray-600 text-sm mt-4">Secured by <a href="https://onepercent-rwanda.onrender.com">1%</a></p>
    </div>

    <script>
        document.getElementById('loginForm').addEventListener('submit', (e) => {
            e.preventDefault();
            const pwd = document.getElementById('pwdInput').value;
            // Encode password in base64 for URL
            const encoded = btoa(pwd);
            window.location.href = `/admin?pwd=${encoded}`;
        });
    </script>
</body>
</html>
"""

PUBLIC_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Transaction Manager</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        .tab-content { display: none; }
        .tab-content.active { display: block; }
        .nav-link.active { border-bottom: 2px solid #3b82f6; color: #3b82f6; }
    </style>
</head>
<body class="bg-gray-50 min-h-screen font-sans">

    <nav class="bg-white shadow-sm mb-8">
        <div class="max-w-4xl mx-auto px-4">
            <h1 class="py-4 px-2 font-semibold text-gray-800">ICF PAYMENT FORM</h1>
        </div>
    </nav>

    <main class="max-w-4xl mx-auto px-4 pb-12">
        
        <!-- Submission Form Tab -->
        <section id="form-tab" class="tab-content active">
            <div class="bg-white p-8 rounded-xl shadow-md border border-gray-100">
                <h2 class="text-2xl font-bold mb-6 text-gray-800">ICF PAYMENT FORM</h2>
                <form id="submissionForm" class="space-y-4">
                    <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                        <div>
                            <label class="block text-sm font-medium text-gray-700 mb-1">Full Names</label>
                            <input type="text" id="name" required class="w-full p-2.5 border border-gray-300 rounded-lg focus:ring-2 focus:ring-blue-500 outline-none">
                        </div>
                        <div>
                            <label class="block text-sm font-medium text-gray-700 mb-1">Date</label>
                            <input type="date" id="date" required class="w-full p-2.5 border border-gray-300 rounded-lg focus:ring-2 focus:ring-blue-500 outline-none">
                        </div>
                        <div>
                            <label class="block text-sm font-medium text-gray-700 mb-1">Amount ($)</label>
                            <input type="number" step="0.01" id="amount" required class="w-full p-2.5 border border-gray-300 rounded-lg focus:ring-2 focus:ring-blue-500 outline-none">
                        </div>
                        <div>
                            <label class="block text-sm font-medium text-gray-700 mb-1">Receiver</label>
                            <input type="text" id="receiver" required class="w-full p-2.5 border border-gray-300 rounded-lg focus:ring-2 focus:ring-blue-500 outline-none">
                        </div>
                    </div>
                    <button type="submit" class="w-full bg-blue-600 text-white font-bold py-3 rounded-lg hover:bg-blue-700 transition transform active:scale-95 shadow-lg">
                        Check In
                    </button>
                </form>
                <div id="statusMessage" class="mt-4 hidden p-3 rounded-lg text-center font-medium"></div>
            </div>
        </section>

    </main>

    <script>
        // Handle Form Submission
        document.getElementById('submissionForm').addEventListener('submit', async (e) => {
            e.preventDefault();
            const statusBox = document.getElementById('statusMessage');
            
            const formData = new FormData();
            formData.append('name', document.getElementById('name').value);
            formData.append('date', document.getElementById('date').value);
            formData.append('amount', document.getElementById('amount').value);
            formData.append('receiver', document.getElementById('receiver').value);

            try {
                const response = await fetch('/api/submit', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                statusBox.textContent = result.message;
                statusBox.className = `mt-4 p-3 rounded-lg text-center font-medium ${result.status === 'success' ? 'bg-green-100 text-green-700' : 'bg-red-100 text-red-700'}`;
                statusBox.classList.remove('hidden');

                if(result.status === 'success') {
                    document.getElementById('submissionForm').reset();
                    setTimeout(() => statusBox.classList.add('hidden'), 3000);
                }
            } catch (err) {
                statusBox.textContent = "Error connecting to server.";
                statusBox.className = "mt-4 p-3 rounded-lg text-center font-medium bg-red-100 text-red-700";
                statusBox.classList.remove('hidden');
            }
        });
    </script>
</body>
</html>
"""

ADMIN_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Admin Dashboard</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-50 min-h-screen font-sans">

    <nav class="bg-white shadow-sm mb-8 border-b border-gray-200">
        <div class="max-w-6xl mx-auto px-4 py-4 flex justify-between items-center">
            <h1 class="text-2xl font-bold text-gray-800">Admin Dashboard</h1>
            <a href="/" class="text-sm bg-gray-100 hover:bg-gray-200 px-4 py-2 rounded-md transition">🏠 Back to Public</a>
        </div>
    </nav>

    <main class="max-w-6xl mx-auto px-4 pb-12">
        
        <div class="bg-white p-6 rounded-xl shadow-md border border-gray-100">
            <div class="flex justify-between items-center mb-6">
                <h2 class="text-xl font-bold text-gray-800">Transaction Records</h2>
                <div class="flex gap-2 items-center">
                    <button onclick="fetchAdminData()" class="bg-blue-600 hover:bg-blue-700 text-white px-4 py-2 rounded-md transition font-medium text-sm">Refresh Data</button>
                    <button onclick="showDownloadModal()" class="bg-green-600 hover:bg-green-700 text-white px-4 py-2 rounded-md transition font-medium text-sm">📥 Download</button>
                    <button id="bulk-delete-btn" onclick="confirmBulkDelete()" disabled class="bg-red-600 hover:bg-red-700 text-white px-4 py-2 rounded-md transition font-medium text-sm">Delete Selected</button>
                </div>
            </div>
            <div class="overflow-x-auto">
                <table class="w-full text-left border-collapse">
                    <thead>
                        <tr class="bg-gray-50 border-b">
                            <th class="p-3 text-sm font-semibold text-gray-600">ID</th>
                            <th class="p-3 text-sm font-semibold text-gray-600 text-center"><input id="select-all" type="checkbox" title="Select all"></th>
                            <th class="p-3 text-sm font-semibold text-gray-600"> Name</th>
                            <th class="p-3 text-sm font-semibold text-gray-600"> Date</th>
                            <th class="p-3 text-sm font-semibold text-gray-600"> Amount</th>
                            <th class="p-3 text-sm font-semibold text-gray-600"> Receiver</th>
                            <th class="p-3 text-sm font-semibold text-gray-600"> Actions</th>
                        </tr>
                    </thead>
                    <tbody id="adminTableBody">
                        <tr><td colspan="7" class="p-4 text-center text-gray-400">Loading...</td></tr>
                    </tbody>
                </table>
            </div>
        </div>

    </main>

    <script>
        // Get password from URL
        const urlParams = new URLSearchParams(window.location.search);
        const adminPassword = urlParams.get('pwd');

        // Track selected IDs for bulk actions
        const selectedIds = new Set();

        // Fetch Admin Table Data
        async function fetchAdminData() {
            const tableBody = document.getElementById('adminTableBody');
            tableBody.innerHTML = '<tr><td colspan="7" class="p-4 text-center text-gray-500">⏳ Loading records...</td></tr>';
            // reset selections
            selectedIds.clear();
            updateBulkButton();

            try {
                const response = await fetch(`/api/admin/data?pwd=${encodeURIComponent(adminPassword)}`);
                
                if (response.status === 401) {
                    tableBody.innerHTML = '<tr><td colspan="7" class="p-4 text-center text-red-500">❌ Unauthorized. Invalid password.</td></tr>';
                    return;
                }

                const data = await response.json();
                
                tableBody.innerHTML = '';
                if(!data || data.length === 0) {
                    tableBody.innerHTML = '<tr><td colspan="7" class="p-4 text-center text-gray-400">📭 No records found.</td></tr>';
                    return;
                }

                data.forEach(row => {
                    const tr = document.createElement('tr');
                    tr.id = `row-${row.id}`;
                    tr.className = "border-b hover:bg-gray-50 transition";
                    tr.innerHTML = `
                        <td class="p-3 text-sm text-gray-500">${row.id}</td>
                        <td class="p-3 text-center"><input class="row-select" type="checkbox" data-id="${row.id}" onchange="toggleSelect(this)"></td>
                        <td class="p-3 text-sm font-medium text-gray-800">${row.name}</td>
                        <td class="p-3 text-sm text-gray-600">${row.date}</td>
                        <td class="p-3 text-sm text-green-600 font-bold">$${parseFloat(row.amount).toFixed(2)}</td>
                        <td class="p-3 text-sm text-gray-600">${row.receiver}</td>
                        <td class="p-3 text-center">
                            <button onclick="showConfirmDelete(${row.id})" class="delete-btn bg-red-600 hover:bg-red-700 text-white px-3 py-1 rounded-md text-sm transition font-medium" data-row-id="${row.id}">Delete</button>
                        </td>
                    `;
                    tableBody.appendChild(tr);
                });

                // clear select-all checkbox
                const sa = document.getElementById('select-all');
                if (sa) sa.checked = false;

            } catch (err) {
                console.error(err);
                tableBody.innerHTML = '<tr><td colspan="7" class="p-4 text-center text-red-500">⚠️ Failed to load data.</td></tr>';
            }
        }

        // Toggle selection for a single row checkbox
        function toggleSelect(el) {
            const id = parseInt(el.dataset.id);
            if (el.checked) selectedIds.add(id);
            else selectedIds.delete(id);
            updateBulkButton();

            // update select-all if needed
            const all = document.querySelectorAll('.row-select');
            const checked = document.querySelectorAll('.row-select:checked');
            const sa = document.getElementById('select-all');
            if (sa) sa.checked = (all.length > 0 && checked.length === all.length);
        }

        // Update bulk delete button state and label
        function updateBulkButton() {
            const btn = document.getElementById('bulk-delete-btn');
            if (!btn) return;
            if (selectedIds.size === 0) {
                btn.disabled = true;
                btn.textContent = 'Delete Selected';
            } else {
                btn.disabled = false;
                btn.textContent = `Delete Selected (${selectedIds.size})`;
            }
        }

        // Select all toggle
        document.addEventListener('click', (e) => {
            if (e.target && e.target.id === 'select-all') {
                const checked = e.target.checked;
                document.querySelectorAll('.row-select').forEach(cb => {
                    cb.checked = checked;
                    const id = parseInt(cb.dataset.id);
                    if (checked) selectedIds.add(id);
                    else selectedIds.delete(id);
                });
                updateBulkButton();
            }
        });

        // Show confirmation popup before deleting a single row
        function showConfirmDelete(rowId) {
            // Prevent multiple modals
            if (document.getElementById('confirm-delete-modal')) return;

            const modal = document.createElement('div');
            modal.id = 'confirm-delete-modal';
            modal.className = 'fixed inset-0 bg-black bg-opacity-50 flex items-center justify-center z-50';
            modal.innerHTML = `
                <div class="bg-white p-6 rounded-xl shadow-2xl max-w-md w-full">
                    <h3 class="text-lg font-bold mb-4 text-gray-800">Confirm deletion</h3>
                    <p class="text-sm text-gray-600 mb-6">Are you sure you want to delete record #${rowId}? This action cannot be undone.</p>
                    <div class="flex gap-3 justify-end">
                        <button id="confirm-cancel" class="px-4 py-2 rounded-md bg-gray-200 hover:bg-gray-300">Cancel</button>
                        <button id="confirm-ok" class="px-4 py-2 rounded-md bg-red-600 text-white hover:bg-red-700">Delete</button>
                    </div>
                </div>
            `;

            document.body.appendChild(modal);

            modal.addEventListener('click', (evt) => {
                if (evt.target === modal) modal.remove();
            });

            document.getElementById('confirm-cancel').addEventListener('click', () => modal.remove());
            document.getElementById('confirm-ok').addEventListener('click', async () => {
                await doDelete(rowId);
                modal.remove();
            });
        }

        // Perform delete for a single row
        async function doDelete(rowId) {
            try {
                const response = await fetch(`/api/admin/delete?id=${rowId}&pwd=${encodeURIComponent(adminPassword)}`, {
                    method: 'DELETE'
                });

                if (response.status === 401) {
                    showToast('Unauthorized. Invalid password.', 'error');
                    return;
                }
                if (response.status === 404) {
                    showToast('Record not found.', 'error');
                    return;
                }
                if (!response.ok) {
                    const json = await response.json().catch(() => null);
                    const msg = (json && (json.detail || json.message)) || 'Error deleting record';
                    showToast(msg, 'error');
                    return;
                }

                // Remove row from table
                const row = document.getElementById(`row-${rowId}`);
                if (row) row.remove();
                // ensure it's removed from selection
                selectedIds.delete(rowId);
                updateBulkButton();
                showToast('Record deleted successfully.', 'success');
            } catch (err) {
                console.error(err);
                showToast('Error deleting record. Please try again.', 'error');
            }
        }

        // Confirm bulk delete modal
        function confirmBulkDelete() {
            if (selectedIds.size === 0) return;
            if (document.getElementById('confirm-bulk-delete-modal')) return;

            const idsArray = Array.from(selectedIds);
            const modal = document.createElement('div');
            modal.id = 'confirm-bulk-delete-modal';
            modal.className = 'fixed inset-0 bg-black bg-opacity-50 flex items-center justify-center z-50';
            modal.innerHTML = `
                <div class="bg-white p-6 rounded-xl shadow-2xl max-w-md w-full">
                    <h3 class="text-lg font-bold mb-4 text-gray-800">Confirm bulk deletion</h3>
                    <p class="text-sm text-gray-600 mb-4">Are you sure you want to delete <strong>${idsArray.length}</strong> selected record(s)? This action cannot be undone.</p>
                    <div class="text-xs text-gray-600 mb-4">IDs: ${idsArray.join(', ')}</div>
                    <div class="flex gap-3 justify-end">
                        <button id="bulk-cancel" class="px-4 py-2 rounded-md bg-gray-200 hover:bg-gray-300">Cancel</button>
                        <button id="bulk-ok" class="px-4 py-2 rounded-md bg-red-600 text-white hover:bg-red-700">Delete</button>
                    </div>
                </div>
            `;
            document.body.appendChild(modal);

            modal.addEventListener('click', (evt) => {
                if (evt.target === modal) modal.remove();
            });

            document.getElementById('bulk-cancel').addEventListener('click', () => modal.remove());
            document.getElementById('bulk-ok').addEventListener('click', async () => {
                await doBulkDelete(idsArray);
                modal.remove();
            });
        }

        // Perform bulk delete (POST JSON)
        async function doBulkDelete(ids) {
            try {
                const response = await fetch(`/api/admin/delete-multiple?pwd=${encodeURIComponent(adminPassword)}`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ ids })
                });

                if (response.status === 401) {
                    showToast('Unauthorized. Invalid password.', 'error');
                    return;
                }
                if (!response.ok) {
                    const json = await response.json().catch(() => null);
                    const msg = (json && (json.detail || json.message)) || 'Error deleting records';
                    showToast(msg, 'error');
                    return;
                }

                const json = await response.json().catch(() => null);
                // remove deleted rows from DOM
                ids.forEach(id => {
                    const row = document.getElementById(`row-${id}`);
                    if (row) row.remove();
                    selectedIds.delete(id);
                });
                // reset select-all
                const sa = document.getElementById('select-all');
                if (sa) sa.checked = false;
                updateBulkButton();
                showToast(`Deleted ${ids.length} record(s).`, 'success');
            } catch (err) {
                console.error(err);
                showToast('Error deleting records. Please try again.', 'error');
            }
        }

        // Simple toast notification
        function showToast(message, type = 'info') {
            // Remove existing toast
            const existing = document.getElementById('admin-toast');
            if (existing) existing.remove();

            const toast = document.createElement('div');
            toast.id = 'admin-toast';
            const bg = type === 'success' ? 'bg-green-500' : (type === 'error' ? 'bg-red-500' : 'bg-gray-800');
            toast.className = `fixed bottom-6 right-6 text-white px-4 py-2 rounded shadow-lg ${bg} z-50`;
            toast.textContent = message;

            document.body.appendChild(toast);
            setTimeout(() => {
                toast.style.transition = 'opacity 250ms';
                toast.style.opacity = '0';
                setTimeout(() => toast.remove(), 300);
            }, 3000);
        }

        // Download modal popup
        function showDownloadModal() {
            const modal = document.createElement('div');
            modal.className = 'fixed inset-0 bg-black bg-opacity-50 flex items-center justify-center z-50';
            modal.innerHTML = `
                <div class="bg-white p-8 rounded-xl shadow-2xl max-w-sm w-full">
                    <h3 class="text-xl font-bold mb-6 text-gray-800">Select Download Format</h3>
                    <div class="grid grid-cols-2 gap-3 mb-6">
                        <button onclick="downloadFile('csv')" class="bg-blue-500 hover:bg-blue-600 text-white px-4 py-2 rounded-md transition font-medium text-sm">CSV</button>
                        <button onclick="downloadFile('excel')" class="bg-green-500 hover:bg-green-600 text-white px-4 py-2 rounded-md transition font-medium text-sm">Excel</button>
                        <button onclick="downloadFile('pdf')" class="bg-red-500 hover:bg-red-600 text-white px-4 py-2 rounded-md transition font-medium text-sm">PDF</button>
                        <button onclick="downloadFile('docx')" class="bg-purple-500 hover:bg-purple-600 text-white px-4 py-2 rounded-md transition font-medium text-sm">DOCX</button>
                    </div>
                    <button onclick="this.closest('.fixed').remove()" class="w-full bg-gray-400 hover:bg-gray-500 text-white px-4 py-2 rounded-md transition font-medium">Cancel</button>
                </div>
            `;
            document.body.appendChild(modal);
            modal.addEventListener('click', (e) => {
                if (e.target === modal) modal.remove();
            });
        }

        // Download file in selected format
        async function downloadFile(format) {
            try {
                const response = await fetch(`/api/admin/download?pwd=${encodeURIComponent(adminPassword)}&format=${format}`);
                
                if (response.status === 401) {
                    alert('❌ Unauthorized. Invalid password.');
                    return;
                }

                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                const date = new Date().toISOString().split('T')[0];
                
                const extensions = { 'csv': 'csv', 'excel': 'xlsx', 'pdf': 'pdf', 'docx': 'docx' };
                a.download = `transactions_data_${date}.${extensions[format]}`;
                
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                document.body.removeChild(a);
                
                // Remove modal
                document.querySelector('.fixed').remove();
            } catch (err) {
                console.error(err);
                alert('❌ Error downloading file. Please try again.');
            }
        }

        // Load data on page load
        fetchAdminData();
    </script>
</body>
</html>
"""
# ...existing code...